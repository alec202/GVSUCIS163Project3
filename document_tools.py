import csv
from docx import Document
import analysis_tools as at

# create DataFile class
class DataFile:
    # define constructor
    def __init__(self, filename, title):
        """
        :param filename: string to represent the name of the file
        :param title: string to represent the title of the file
        """
        self.filename = filename
        self.title = title
        self.description = None
        self.process = []
        self.data = []
        self.header = None

    """beginning of setters and getters"""
    def set_filename(self, filename):
        """
        :param filename: string to represent the new filename entered
        :return: nothing
        """
        self.filename = filename

    def get_filename(self):
        """
        :return: the filename for the object
        """
        return self.filename

    def set_title(self, title):
        """
        :param title: string to represent the title entered
        :return: nothing
        """
        self.title = title

    def get_title(self):
        """
        :return: the title of th object
        """
        return self.title

    def set_description(self, description):
        """
        :param description: string to represent the description of the object
        :return: nothing
        """
        self.description = description

    def get_description(self):
        """
        :return: the description of the associated object
        """
        return self.description

    def set_process(self, process):
        """
        :param process: list of all the columns we wish to process
        :return: nothing
        """
        self.process = process

    def get_process(self):
        """
        :return: the list of all the columns we will process
        """
        return self.process

    def set_data(self, data):
        """
        :param data: list that will hold the data of the file
        :return: nothing
        """
        self.data = data

    def get_data(self):
        """
        :return: the list of the data of the file
        """
        return self.data

    def set_header(self, header):
        """
        :param header: list of the names of the columns
        :return: nothing
        """
        self.header = header

    def get_header(self):
        """
        :return: the list of the names of the columns
        """
        return self.header

    """End of setters and getters"""

    def get_col_name(self, index):
        """
        :param index: index to determine which name of a specified column is sought
        :return: the name of the specified column at index
        """
        return self.header[index]

    def get_col(self, col):
        """
        :param col: the column that we want data for
        :return: a new list with only the data of the column specified
        """
        new_list = []
        for row in self.data:
            new_list.append(row[col])
        return new_list

    def load(self):
        """
        :return: nothing
        """
        data = []
        with open(self.filename, encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                data.append(row)
            self.header = data[0]
            del (data[0])
            self.data = data

# create ReportGenerator class
class ReportGenerator:
    # definite the constructor
    def __init__(self, datafile):
        """
        :param datafile: an instance of DataFile class that was passed in
        """
        self.datafile = datafile

    # define function to actually make and save the document
    def generate(self, output_name):
        """
        :param output_name: name to save the document with
        :return:nothing
        """

        # creating the actual report document
        document = Document()
        document.add_heading('heading', 0)
        document.add_paragraph('description')

        # DataFile.get_title()
        for col in self.datafile.get_process():
            document.add_heading(self.datafile.get_col_name(col), 1)
            # creating the table
            table = document.add_table(rows=1, cols=2)

            # all the rows in the table and their contents being changed/corrected
            row_cells1 = table.add_row().cells
            row_cells1[0].text = "Mean"
            mean = at.MyMath.mean(self.datafile.get_col(col))
            row_cells1[1].text = str(mean)

            row_cells2 = table.add_row().cells
            row_cells2[0].text = "Standard Deviation"
            stDev = at.MyMath.stdev(self.datafile.get_col(col))
            row_cells2[1].text = str(stDev)

            row_cells3 = table.add_row().cells
            row_cells3[0].text = "Median"
            mean = at.MyMath.mean(self.datafile.get_col(col))
            row_cells3[1].text = str(mean)

            row_cells4 = table.add_row().cells
            row_cells4[0].text = "Minimum"
            min = at.MyMath.minimum(self.datafile.get_col(col))
            row_cells4[1].text = str(min)

            row_cells5 = table.add_row().cells
            row_cells5[0].text = "Maximum"
            max = at.MyMath.maximum(self.datafile.get_col(col))
            row_cells5[1].text = str(max)

        document.save(output_name)

'''for col in self.datafile.get_col_name():
table = document.add_table(rows=1, cols=2)
row_cells = table.add_row().cells
row_cells[0].text = "Mean"
row_cells[1].text = at.MyMath.mean(DataFile.get_col(col)) + str(42) 
document.save(output_name)''
